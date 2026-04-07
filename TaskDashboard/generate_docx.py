import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Portada
title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("\n\n\n\n\n\nAntigraviti\nDocumentación Técnica del Proyecto")
title_run.bold = True
title_run.font.size = Pt(28)

document.add_page_break()

# Tabla de Contenidos - Placeholder simple
doc_title = document.add_heading("Tabla de Contenido", level=1)
document.add_paragraph("1. 📁 Función de las Carpetas")
document.add_paragraph("2. 🎨 Íconos y Logos")
document.add_paragraph("3. 🔑 API")
document.add_paragraph("4. ✍️ Firma Digital")
document.add_paragraph("5. 📦 AAB y APK")
document.add_paragraph("6. 🐙 GitHub")

document.add_page_break()

# 1
document.add_heading("1. 📁 FUNCIÓN DE LAS CARPETAS", level=1)
document.add_paragraph("NOTA IMPORTANTE: El proyecto está construido nativamente bajo React Native con TypeScript. La estructura del proyecto obedece a React Native, la cual es la herramienta predominante con la que fue generada la aplicación.")
document.add_paragraph("src/: Aquí está todo el código principal de la aplicación. Contiene la interfaz, componentes, navegación y la lógica TypeScript que gobierna la app.", style='List Bullet')
document.add_paragraph("android/: Contiene el proyecto nativo de Android. Adentro está todo lo necesario para compilar el APK/AAB y las configuraciones de Gradle.", style='List Bullet')
document.add_paragraph("ios/: Contiene el proyecto nativo de iOS (Xcode). Funciona igual que la carpeta de Android pero para Apple.", style='List Bullet')
document.add_paragraph("node_modules/: Aquí se instalan todas las librerías y dependencias externas del proyecto (generada automáticamente, no se sube a GitHub ni se edita a mano).", style='List Bullet')
document.add_paragraph("__tests__/: Utilizada para guardar las pruebas automáticas o unitarias del proyecto.", style='List Bullet')
document.add_paragraph("assets/: Usualmente almacenada en src/assets (o configuraciones externas), alberga imágenes o recursos gráficos estáticos.", style='List Bullet')

# 2
document.add_heading("2. 🎨 ÍCONOS Y LOGOS", level=1)
document.add_heading("Ruta exacta:", level=2)
p = document.add_paragraph("• Android: ")
p.add_run("android/app/src/main/res/mipmap-*").font.name = 'Courier New'
p = document.add_paragraph("• iOS: ")
p.add_run("ios/TaskDashboard/Images.xcassets/AppIcon.appiconset/").font.name = 'Courier New'

document.add_heading("Cómo y dónde modificarlos:", level=2)
document.add_paragraph("Existe un archivo generado llamado generate_icons.py en la raíz del proyecto. Si colocas un nuevo logo base y ejecutas el script, este automáticamente procesará las resoluciones y reemplazará todos los tamaños necesarios en Android e iOS de una sola vez.")

document.add_heading("Dónde está el logo principal y cómo reemplazarlo:", level=2)
document.add_paragraph("A través de dicho script Python (generate_icons.py) usando el logo de origen. De forma manual deberás reemplazar cada imagen .png en las resoluciones respectivas (hdpi, mdpi, xhdpi, xxhdpi) ubicadas en las rutas nativas mencionadas arriba.")

# 3
document.add_heading("3. 🔑 API", level=1)
document.add_heading("Archivo de configuración:", level=2)
p = document.add_paragraph()
p.add_run("src/data/datasources/RemoteDataSource.ts").font.name = 'Courier New'

document.add_heading("URLs y Keys:", level=2)
document.add_paragraph("En ese archivo, la instancia de Axios está configurada de la siguiente forma:")
code = document.add_paragraph()
code.add_run("const api: AxiosInstance = axios.create({\n  baseURL: 'https://dummyjson.com',\n  timeout: 10_000,\n  headers: { 'Content-Type': 'application/json' },\n});").font.name = 'Courier New'
document.add_paragraph("Aquí es donde debes modificar baseURL ('https://dummyjson.com') por tu servidor y base de datos oficial. Adicionalmente, las Keys de Autenticación de Firebase u otras plataformas suelen configurarse en archivos .env (dotenv).")

# 4
document.add_heading("4. ✍️ FIRMA DIGITAL", level=1)
document.add_heading("Ubicación del Archivo:", level=2)
p = document.add_paragraph()
p.add_run("android/app/taskdashboard-release.jks").font.name = 'Courier New'

document.add_heading("Datos de la Firma:", level=2)
document.add_paragraph("• Nombre del archivo: taskdashboard-release.jks")
document.add_paragraph("• Alias de la llave (Key alias): taskdashboard")
document.add_paragraph("• Contraseña del Keystore: TaskDashboard2024!")
document.add_paragraph("• Contraseña de la Llave (Key password): TaskDashboard2024!")
document.add_paragraph("Nota: Estas contraseñas son requeridas por Gradle para la construcción y se asignaron en android/gradle.properties.")

# 5
document.add_heading("5. 📦 AAB y APK", level=1)
document.add_heading("Dónde se generan:", level=2)
p = document.add_paragraph("• APK: ")
p.add_run("android/app/build/outputs/apk/release/app-release.apk").font.name = 'Courier New'
p = document.add_paragraph("• AAB: ")
p.add_run("android/app/build/outputs/bundle/release/app-release.aab").font.name = 'Courier New'

document.add_heading("Comandos para compilarlos:", level=2)
document.add_paragraph("Abre la consola CMD o PowerShell en la carpeta del proyecto y ejecuta:")
c1 = document.add_paragraph("Para APK (instalación directa):")
c1.add_run("\ncd android\n.\\gradlew assembleRelease").font.name = 'Courier New'

c2 = document.add_paragraph("Para AAB (Publicación en Google Play Store):")
c2.add_run("\ncd android\n.\\gradlew bundleRelease").font.name = 'Courier New'

# 6
document.add_heading("6. 🐙 GITHUB", level=1)
document.add_heading("Paso a paso para subir el proyecto:", level=2)
document.add_paragraph("1. Inicializa Git en tu directorio.\n2. Asegúrate de añadir las credenciales a .gitignore para omitirlas.\n3. Añade todos los archivos restantes.\n4. Crea un \"commit\" (punto de restauración).\n5. Cambia el nombre de la rama a main.\n6. Enlaza con tu URL del repositorio de GitHub.\n7. Sube las modificaciones por primera vez (push).")

document.add_heading("Comandos exactos de Git:", level=2)
git_cmd = document.add_paragraph()
git_cmd.add_run("git init\ngit add .\ngit commit -m \"Primer commit de la aplicación\"\ngit branch -M main\ngit remote add origin <TU_URL_DE_GITHUB>\ngit push -u origin main").font.name = 'Courier New'

document.add_heading("Qué archivos NO subir y por qué:", level=2)
document.add_paragraph("• *.jks (Tu firma digital): Es el certificado digital. Quien tenga esto y las contraseñas podrá subir copias maliciosas pasándolas como tuyas en la tienda.")
document.add_paragraph("• android/gradle.properties: O al menos remueve la sección de passwords. No quieres que el mundo vea 'TaskDashboard2024!'.")
document.add_paragraph("• node_modules / build / .gradle / .idea: Son carpetas de archivos transitorios, cachés o librerías descargables. Son pesadísimas, se generan solas al compilar y si las subes arruinarás el tamaño de tu repositorio y generarás conflictos a nivel código en otros equipos.")

# Guardar documento
document.save("c:/Users/ASUS/OneDrive/Desktop/APP-GM/Documentacion_Antigraviti.docx")
print("Documento guardado exitosamente en APP-GM")
